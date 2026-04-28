"""
Build SwarmGen_Runbook.docx from a single Python script. No markdown, no pandoc.
Run:  python build_runbook.py
Writes: SwarmGen_Runbook.docx in the project root.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Mm, Inches


OUT = Path(__file__).parent / "SwarmGen_Runbook.docx"


# ---------- helpers ---------------------------------------------------------
def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="C8C8C8", sz="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def code_block(doc: Document, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, "F4F4F5")
    set_cell_borders(cell, color="E4E4E7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x23)
    # spacer paragraph
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def inline_code(p, text: str):
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x21, 0x80)
    return run


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)


def make_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            tbl.rows[r_i].cells[c_i].text = ""
            p = tbl.rows[r_i].cells[c_i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()


# ---------- build ------------------------------------------------------------
def build() -> Document:
    doc = Document()

    # Page margins, smaller.
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===========================================================================
    # COVER
    # ===========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("SwarmGen")
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)

    sub = doc.add_paragraph()
    run = sub.add_run("Distributed Stable Diffusion Turbo across heterogeneous edge devices.")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

    meta = doc.add_paragraph()
    run = meta.add_run("CIS 589 Edge Computing  ·  Final Project  ·  Spring 2026  ·  University of Michigan-Dearborn")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    meta2 = doc.add_paragraph()
    run = meta2.add_run("Sudarshan Sridhar  ·  Varun Patel  ·  github.com/sudarshan-sridhar/swarmgen")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "This is the operational runbook. It explains what the project is, why we built it the "
        "way we did, how to bring it up cold from three SSH sessions, what the UI shows, and "
        "what to say while recording the demo. Read it once before going live. Keep it open while "
        "running."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ===========================================================================
    # 1. WHAT THIS IS
    # ===========================================================================
    add_heading(doc, "1. What this is, in one paragraph", 1)
    p = doc.add_paragraph(
        "SwarmGen takes Stable Diffusion Turbo, a four-step text-to-image diffusion model, and "
        "splits its three components across three different edge devices on the same Wi-Fi. "
        "The text encoder runs on a CPU laptop. The UNet denoiser runs on a laptop with an "
        "RTX 5060. The VAE decoder runs on a Raspberry Pi 4B. A Python coordinator on the GPU "
        "laptop ties them together over async HTTP, watches their heartbeats, and reroutes the "
        "VAE stage to the GPU when the Pi dies mid-flight. There is a small static web UI for "
        "the live demo."
    )

    p = doc.add_paragraph(
        "The point is not speed. The point is that the Pi physically cannot run the full "
        "pipeline on its own, the swarm makes it a participant, and the system survives a "
        "device dying without dropping the image."
    )

    # ===========================================================================
    # 2. THE PROBLEM
    # ===========================================================================
    add_heading(doc, "2. The problem", 1)
    doc.add_paragraph(
        "Diffusion image models like Stable Diffusion sit just outside the envelope of small "
        "edge devices. The smallest practical version, SD-Turbo, has a peak working set of "
        "about 2.4 GB across CPU and GPU when you load all components together. A Raspberry Pi "
        "4B with 1.8 GB of RAM cannot fit that. It is not a software problem you can fix with "
        "swap. It is a hard ceiling."
    )
    doc.add_paragraph(
        "The usual answer is to put the model on a beefy machine and call it from the edge. "
        "That works, but it concedes the point. The interesting question is whether a swarm "
        "of unequal devices, each holding a piece of the model, can do something none of them "
        "could do alone, and whether the swarm degrades gracefully when a device fails."
    )
    doc.add_paragraph(
        "This is the question we set out to answer. Not to beat a single GPU on per-image "
        "latency, because we will lose that fight by two orders of magnitude. The honest "
        "claim is memory partitioning, end-to-end correctness, and fault tolerance."
    )

    # ===========================================================================
    # 3. THE SOLUTION
    # ===========================================================================
    add_heading(doc, "3. The solution", 1)
    doc.add_paragraph(
        "Split the pipeline into three stages, one per device, each holding only the model "
        "component it can run."
    )

    add_heading(doc, "3.1 Pipeline split", 2)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("CLIP text encoder, about 120 M parameters, on the CPU laptop. One forward pass "
              "per prompt. CPU is fine.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("UNet denoiser, about 865 M parameters, on the RTX 5060. Four forward passes per "
              "image. Has to be on the GPU to be tractable.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("VAE decoder, about 80 M parameters, on the Pi. One forward pass per image. Small "
              "enough to fit in 1.8 GB of RAM, slow at about 140 seconds per 512 by 512 decode "
              "in CPU FP32. Slow but bounded.")

    doc.add_paragraph(
        "Putting UNet on the GPU is not a cheat. The paper claim is heterogeneity and per-device "
        "memory reduction, not that the GPU does less work. We measured peak RSS on every device. "
        "The Pi peaks at about 1225 MB, which fits in its 1.8 GB. The single-device baseline "
        "peaks at about 2358 MB, which does not. That is the headline."
    )

    add_heading(doc, "3.2 Discovery and orchestration", 2)
    doc.add_paragraph(
        "Workers announce themselves over mDNS under the service name "
    ).add_run("_swarmgen._tcp.local.").font.name = "Consolas"
    doc.add_paragraph(
        "The coordinator finds them. There is also an explicit "
    ).add_run("--workers role@host:port,...").font.name = "Consolas"
    doc.add_paragraph(
        "flag for the times mDNS is flaky. We lean on the explicit flag during the demo because "
        "it is faster and more predictable on stage."
    )

    add_heading(doc, "3.3 Heartbeat and fault recovery", 2)
    doc.add_paragraph(
        "A background async task in the coordinator polls every worker's /heartbeat endpoint "
        "once a second. After three misses in a row, the worker is marked dead and excluded "
        "from the candidate pool for new stage calls. Stage calls themselves catch transport "
        "errors immediately, so the next candidate is tried without waiting for the heartbeat "
        "to confirm."
    )
    doc.add_paragraph(
        "Workers can declare extra roles they support via --fallback-roles. The GPU laptop "
        "runs UNet primary plus VAE fallback. Adding the VAE component to the GPU costs about "
        "165 MB of VRAM, well under our 8 GB budget. When the Pi dies, the coordinator retries "
        "VAE on the GPU and the image still completes, in roughly 250 ms instead of 140 seconds."
    )

    add_heading(doc, "3.4 The asyncio inference fix", 2)
    doc.add_paragraph(
        "There is a subtle bug we hit and fixed. PyTorch inference is synchronous CPU work. If "
        "the worker's /run_stage runs torch directly inside an async FastAPI endpoint, the "
        "event loop is pinned for the entire decode. The Pi's 140-second VAE decode ends up "
        "blocking /heartbeat too, so the coordinator falsely declares the Pi dead a few seconds "
        "in. We push the model call into asyncio.to_thread so the loop stays free for health "
        "and heartbeat probes during inference. This was the difference between the heartbeat "
        "monitor working and not."
    )

    # ===========================================================================
    # 4. DEVICES
    # ===========================================================================
    add_heading(doc, "4. The three devices", 1)
    doc.add_paragraph(
        "We use the short names loq, pc, and pi everywhere. They appear in logs, the UI, the "
        "paper, and the rest of this document. Locked specs:"
    )

    make_table(doc,
        ["short", "long hostname", "role", "specs", "Wi-Fi IP", "port"],
        [
            ["loq", "SudarshanLOQ",        "UNet (+ VAE fallback)", "Win 11, RTX 5060 Laptop 8 GB VRAM, 32 GB RAM, 20 logical CPUs", "192.168.1.16",  "8002"],
            ["pc",  "Sudarshan-PC",        "CLIP",                  "Win 11, Intel i5-1035G1, 12 GB RAM, no GPU",                    "192.168.1.39",  "8001"],
            ["pi",  "raspberrypiinocula",  "VAE",                   "Pi 4B, ARM Cortex-A72, 1.8 GB RAM, Debian 13, no GPU",         "192.168.1.185", "8003"],
        ]
    )

    doc.add_paragraph(
        "All three sit on the same home Wi-Fi (192.168.1.0/24). loq also runs the coordinator "
        "and the web UI."
    )

    # ===========================================================================
    # 5. WHAT'S IN THE REPO
    # ===========================================================================
    add_heading(doc, "5. What's in the repo", 1)
    make_table(doc,
        ["file", "what it does"],
        [
            ["protocol.py",       "Tensor wire format. 4-byte length prefix, JSON header, raw bytes. No pickle."],
            ["worker.py",         "Single worker. Role-flagged. Loads only its assigned model component plus any fallback roles. Exposes /health, /heartbeat, /capabilities, /run_stage, /admin/die."],
            ["coordinator.py",    "Async orchestrator. mDNS discovery and direct --workers. Single image and pipeline-parallel batch. Heartbeat monitor and fallback retry."],
            ["api.py",            "FastAPI server in front of the coordinator. Serves the static UI plus a small JSON API: /api/workers, /api/generate, /api/admin/kill-vae."],
            ["static/index.html", "Single-file dark-themed web UI. Vanilla HTML, Tailwind CDN, Bricolage Grotesque + JetBrains Mono. No build step."],
            ["baseline.py",       "Single-device SD-Turbo baseline on loq. Used for the latency comparison."],
            ["eval.py",           "Eval harness. Subcommands latency, memory, network, batch, fault. Writes results CSVs."],
            ["plot_results.py",   "Reads results/*.csv and writes the four paper figures into paper/figs/."],
            ["paper/paper.tex",   "IEEE conference paper, eight sections, Overleaf-ready."],
            ["demo_script.md",    "Beat-by-beat 7 minute demo script. (This document is the longer companion to it.)"],
            ["DESIGN.md",         "The design system the UI was built against."],
            ["CONTEXT.md",        "Original project context handed to Claude Code."],
        ]
    )

    # ===========================================================================
    # 6. COLD-START RUNBOOK
    # ===========================================================================
    add_heading(doc, "6. Cold-start runbook (you are at loq, three SSH tabs)", 1)
    doc.add_paragraph(
        "This section assumes you walk up to loq with nothing running, and you want all three "
        "workers up plus the UI in front of you, all driven from this one machine. The pi and "
        "pc are powered on and on the same Wi-Fi but no software is running yet."
    )
    doc.add_paragraph(
        "Open three terminal windows on loq. One stays local. The other two SSH into pc and pi."
    )

    add_heading(doc, "Tab 1 — local on loq (UNet + coordinator + UI)", 2)
    doc.add_paragraph(
        "This is a regular PowerShell window on loq, conda env ml. We start two processes: "
        "the UNet worker, and the API server that fronts the coordinator and serves the UI."
    )
    code_block(doc,
        "conda activate ml\n"
        "cd C:\\Users\\sudar\\projects\\school\\swarmgen\n"
        "git pull\n"
        "\n"
        "# 1) UNet worker. Loads UNet plus VAE-fallback on the GPU.\n"
        "python worker.py --role unet --port 8002 --fallback-roles vae"
    )
    doc.add_paragraph(
        "Wait until the worker logs the banner SwarmGen worker UP role=unet. First boot has "
        "to download the SD-Turbo weights (about 3.4 GB) into the Hugging Face cache. After "
        "the first time it loads in 30 to 90 seconds. Leave this window open. The worker logs "
        "every stage call, which is useful to watch live."
    )
    doc.add_paragraph("In a second tab on loq:")
    code_block(doc,
        "conda activate ml\n"
        "cd C:\\Users\\sudar\\projects\\school\\swarmgen\n"
        "\n"
        "# 2) API server + UI. This is what your browser hits.\n"
        "python api.py --workers \"clip@192.168.1.39:8001,unet@192.168.1.16:8002,vae@192.168.1.185:8003\""
    )
    doc.add_paragraph(
        "It prints a line like SwarmGen UI ready at http://localhost:7860. Open that in a "
        "browser. The page itself is one HTML file in static/ so you can edit it and just "
        "refresh the tab."
    )

    add_heading(doc, "Tab 2 — SSH into pc (CLIP)", 2)
    doc.add_paragraph(
        "From loq, SSH into pc. The username has a space in it, so quote it. Tailscale also "
        "exposes pc at 100.68.82.88 if your Wi-Fi is being weird."
    )
    code_block(doc,
        "# from loq PowerShell\n"
        "ssh \"Sudarshan Sridhar@100.68.82.88\"\n"
        "# (or: ssh \"Sudarshan Sridhar@192.168.1.39\")"
    )
    doc.add_paragraph("Once the cmd prompt for pc shows up, run:")
    code_block(doc,
        "cd swarmgen\n"
        "git pull\n"
        ".venv\\Scripts\\activate.bat\n"
        "python worker.py --role clip --port 8001"
    )
    doc.add_paragraph(
        "CLIP is small. About 280 MB RSS, 30 to 60 seconds to load. Wait for the worker UP "
        "banner. Leave the SSH session open."
    )

    add_heading(doc, "Tab 3 — SSH into pi (VAE)", 2)
    doc.add_paragraph(
        "From loq, SSH into the Pi. Hostname is raspberrypiinocula. mDNS works on the LAN."
    )
    code_block(doc,
        "ssh pi@raspberrypiinocula.local\n"
        "# fallback by IP if mDNS is flaky:\n"
        "# ssh pi@192.168.1.185"
    )
    doc.add_paragraph("Once you are at the pi prompt:")
    code_block(doc,
        "cd ~/swarmgen\n"
        "git pull\n"
        "source .venv/bin/activate\n"
        "python worker.py --role vae --port 8003"
    )
    doc.add_paragraph(
        "The Pi takes longer because it is the slowest box. About 90 seconds to import torch "
        "and load the VAE. Watch for SwarmGen worker UP role=vae. Leave the SSH session open."
    )

    add_heading(doc, "Sanity check before going live", 2)
    doc.add_paragraph(
        "From loq, in any spare terminal, hit each /health endpoint. All three should return "
        "JSON with status: ok."
    )
    code_block(doc,
        "curl http://192.168.1.16:8002/health   # loq UNet\n"
        "curl http://192.168.1.39:8001/health   # pc CLIP\n"
        "curl http://192.168.1.185:8003/health  # pi VAE"
    )
    doc.add_paragraph(
        "If any of them errors, do not move on. Fix it before the demo starts. See the "
        "troubleshooting section."
    )
    doc.add_paragraph(
        "Now refresh http://localhost:7860 on loq. The workers table at the bottom should show "
        "three rows, all with the green pulsing alive pill."
    )

    # ===========================================================================
    # 7. WHAT THE UI SHOWS
    # ===========================================================================
    add_heading(doc, "7. What the UI shows, section by section", 1)
    doc.add_paragraph(
        "The page is one HTML file at static/index.html. It is dark themed, uses Bricolage "
        "Grotesque for type and JetBrains Mono for everything numeric. Tailwind is loaded from "
        "the CDN at runtime, no build step. The design system that drives it is in DESIGN.md "
        "in the repo root."
    )

    add_heading(doc, "Hero", 2)
    doc.add_paragraph(
        "Left aligned title, italic accent on the word edges, a short description and a thin "
        "metabar with model, steps, output size, devices, protocol. No marketing fluff."
    )

    add_heading(doc, "01 / control", 2)
    doc.add_paragraph(
        "Prompt box, three small inputs (seed, steps, fault inject), and two buttons. Generate "
        "is the emerald primary action. Kill Pi VAE worker is the restrained crimson danger "
        "button. The summary line under the buttons reports either Total xxx ms after a "
        "successful run, or Failed plus the reason if anything went wrong. Cmd or Ctrl + Return "
        "in the prompt box also fires generate."
    )

    add_heading(doc, "02 / pipeline", 2)
    doc.add_paragraph(
        "Per-stage timing table that fills in after each generation. CLIP, UNET, VAE on the "
        "left. ms and bytes on the right, in monospace tabular numerics. The ms column has a "
        "tiny emerald heat-bar that scales relative to the slowest stage in this run, so you "
        "can see at a glance which stage dominated. The via column shows the host and port of "
        "the worker that actually served the stage. The retry count is on the right end."
    )
    doc.add_paragraph(
        "When the Pi is alive and you generate, VAE shows up as the dominant bar at about 140 s. "
        "When the Pi dies and the fallback fires, VAE shows up as a tiny bar at a couple "
        "hundred ms, with the via column pointing at loq instead of the Pi."
    )

    add_heading(doc, "03 / output", 2)
    doc.add_paragraph(
        "The image frame. Dark surface, hairline border, 16 px radius. Empty state shows a "
        "small ASCII trace of the pipeline so the page is not blank on first load. While the "
        "generation is running, a slow shimmer moves across the frame to make it obvious "
        "something is happening. After the run, the new image fades in. The status chip in "
        "the section header switches between idle, running, ok, and fail."
    )

    add_heading(doc, "04 / swarm", 2)
    doc.add_paragraph(
        "The worker telemetry table. Polls /api/workers every two seconds. Each row has a "
        "pulsing green dot next to the role tag for live workers and a static red dot for "
        "dead ones. The status pill confirms that. The memory column shows current RSS over "
        "peak in monospace numbers, with a thin emerald bar underneath that uses the largest "
        "peak across all workers as the scale, so the relative memory pressure across loq, pc, "
        "and pi is visible at a glance."
    )

    # ===========================================================================
    # 8. DEMO SCRIPT
    # ===========================================================================
    add_heading(doc, "8. Demo recording script (target 7 minutes)", 1)
    doc.add_paragraph(
        "This is what to actually say while recording. Tone is casual and technical. Read it "
        "out loud once before recording so you do not stumble on the transitions. No reading "
        "the slides word for word. No marketing language."
    )
    doc.add_paragraph(
        "Setup before you hit record: all three workers running, the UI tab open at "
        "http://localhost:7860, a second tab or window with the loq worker log visible, and "
        "the architecture figure or memory plot open in a third tab for the visuals."
    )

    add_heading(doc, "0:00 to 0:45  the pitch", 2)
    doc.add_paragraph(
        "On screen: the UI tab. Speak over it."
    )
    doc.add_paragraph(
        "What to say, roughly: \"This is SwarmGen, our final project for CIS 589 at "
        "UM-Dearborn. The idea is to take Stable Diffusion Turbo, a four-step diffusion image "
        "model, and split it across three edge devices on the same Wi-Fi. The text encoder "
        "runs on a CPU laptop, the UNet runs on a laptop with an RTX 5060, the VAE runs on a "
        "Raspberry Pi 4B. Why bother. Because the Pi has 1.8 gigabytes of RAM and the full "
        "SD-Turbo pipeline needs 2.4. The Pi physically cannot run this on its own. SwarmGen "
        "lets it participate.\""
    )

    add_heading(doc, "0:45 to 1:30  what we are not doing", 2)
    doc.add_paragraph(
        "On screen: the latency comparison plot from paper/figs/latency_comparison.png."
    )
    doc.add_paragraph(
        "What to say: \"I want to be honest up front. We are not trying to beat a single GPU "
        "on per-image latency. Putting a Pi in the pipeline adds 140 seconds of VAE decode that "
        "the GPU would have done in a quarter of a second. The point of the project is memory "
        "partitioning and fault tolerance, not raw speed. The single-device baseline on the "
        "GPU laptop runs an image in about 380 milliseconds. The three-device swarm with the "
        "Pi running VAE takes 147 seconds. We will report what we measured.\""
    )

    add_heading(doc, "1:30 to 3:30  live single-image demo", 2)
    doc.add_paragraph(
        "On screen: the UI tab. Worker telemetry table at the bottom should show three "
        "alive rows."
    )
    doc.add_paragraph(
        "Type a prompt. Suggested: a watercolor painting of a lighthouse at dusk. Leave fault "
        "inject as none. Click generate."
    )
    doc.add_paragraph(
        "What to say while it runs: \"You can see the per-stage timing fill in as the pipeline "
        "moves through the workers. CLIP ran on the CPU laptop in about 500 milliseconds, you "
        "can see the via column pointing at 192.168.1.39 port 8001. UNet ran on the GPU in "
        "230 milliseconds. Now we are sitting on VAE on the Pi, this is the slow one, going to "
        "take about two and a half minutes. The shimmer in the image frame tells you we are "
        "still working. While we wait, take a look at the worker telemetry at the bottom. The "
        "green pulse is heartbeat. The Pi memory is climbing as the decode progresses.\""
    )
    doc.add_paragraph(
        "When the image lands, point at the per-stage timing again. \"There it is. The VAE "
        "row takes about 140 seconds and dwarfs the rest, exactly as we said. The image is "
        "real. You can see the path of the file in the meta line under the frame.\""
    )

    add_heading(doc, "3:30 to 5:00  fault tolerance", 2)
    doc.add_paragraph(
        "This is the headline of the demo. Set the fault inject dropdown to vae. Type a "
        "different prompt. Suggested: a robot fox in a cyberpunk city. Click generate."
    )
    doc.add_paragraph(
        "What to say: \"Now the fun part. The GPU laptop is configured with VAE as a fallback "
        "role, so it has the VAE model loaded too, costs about 165 megabytes of VRAM, totally "
        "trivial on the 8 gig card. The fault inject I just selected schedules a kill of the "
        "Pi worker one second into the generation. Watch what happens.\""
    )
    doc.add_paragraph(
        "While it runs, the heartbeat monitor in the coordinator will catch the Pi going dead, "
        "and the next stage call to VAE will fail at the transport layer and immediately retry "
        "on the GPU laptop's fallback. The image renders in roughly 5 to 6 seconds total."
    )
    doc.add_paragraph(
        "What to say after: \"The image rendered in about five seconds. Look at the per-stage "
        "table. VAE ran on 192.168.1.16 port 8002. That is loq. The Pi is gone. The retry "
        "count column shows the failed attempt against the Pi. The worker telemetry now shows "
        "the Pi as DEAD. The image still made it.\""
    )

    add_heading(doc, "5:00 to 6:00  batch throughput", 2)
    doc.add_paragraph(
        "On screen: the throughput plot from paper/figs/throughput_scaling.png."
    )
    doc.add_paragraph(
        "What to say: \"There is also a batch mode in the coordinator with three async tasks "
        "and bounded queues between them. While CLIP encodes prompt three, UNet denoises "
        "prompt two, VAE decodes prompt one. Pipeline parallelism. When the GPU runs VAE, "
        "throughput is 87 images per minute. When the Pi runs VAE, throughput is 0.41 images "
        "per minute. The reason is steady-state throughput is bounded by the slowest stage, "
        "and the Pi is two orders of magnitude slower. Pipeline parallelism cannot save you "
        "when the stages are this unbalanced. That is one of the conclusions in the paper.\""
    )

    add_heading(doc, "6:00 to 6:45  memory reduction", 2)
    doc.add_paragraph(
        "On screen: paper/figs/memory_per_device.png."
    )
    doc.add_paragraph(
        "What to say: \"And here is the memory story, which is the actual reason this project "
        "exists. The single-device baseline peaks at 2358 megabytes. The Pi peaks at 1226, "
        "well under its 1.8 gigabyte ceiling. The dashed red line is the Pi's physical RAM. "
        "The single-device peak is above it. The Pi cannot run the full pipeline alone. The "
        "swarm makes it a participant. That is the headline.\""
    )

    add_heading(doc, "6:45 to 7:00  wrap", 2)
    doc.add_paragraph(
        "What to say: \"Source code, paper, and result CSVs are at "
        "github.com/sudarshan-sridhar/swarmgen. About 1500 lines of Python, plus the static "
        "frontend. FastAPI, asyncio, zeroconf, Tailwind from the CDN. No Docker, no Ray, no "
        "Kubernetes. Thanks for watching.\""
    )

    # ===========================================================================
    # 9. EVAL NUMBERS
    # ===========================================================================
    add_heading(doc, "9. Eval numbers worth memorizing", 1)
    doc.add_paragraph(
        "These are the actual numbers from results/ and outputs/. Memorize them so you can "
        "speak them naturally during the recording."
    )
    make_table(doc,
        ["metric", "value", "where it lives"],
        [
            ["1-device baseline mean latency",       "380 ms steady (480 ms incl. warm-up)", "results/baseline_latency.csv"],
            ["3-device swarm with loq VAE fallback", "1097 ms per image",                    "results/latency_3dev_loq_vae.csv"],
            ["3-device swarm with Pi VAE primary",   "147 615 ms per image",                 "results/latency_3dev_pi.csv"],
            ["1-device peak RSS",                    "2358 MB",                              "results/baseline_latency.csv"],
            ["Pi VAE peak RSS",                      "1226 MB (under 1845 MB ceiling)",      "results/memory_3dev_pi.csv"],
            ["loq UNet peak RSS",                    "2080 MB",                              "results/memory_3dev_pi.csv"],
            ["pc CLIP peak RSS",                     "1398 MB",                              "results/memory_3dev_pi.csv"],
            ["Batch throughput, loq VAE",            "87.6 img/min",                         "outputs/batch_no_pi/batch_summary.json"],
            ["Batch throughput, Pi VAE",             "0.41 img/min",                         "outputs/batch_with_pi/batch_summary.json"],
            ["Fault recovery, Pi to loq VAE",        "about 250 ms",                         "outputs/a-robot-fox-...__fault-vae.json"],
        ]
    )

    # ===========================================================================
    # 10. UI ↔ ARCHITECTURE MAP
    # ===========================================================================
    add_heading(doc, "10. How the UI matches the architecture", 1)
    doc.add_paragraph(
        "Quick mental map for when you talk over the UI on camera. Every visible piece of the "
        "page corresponds to a real piece of the system."
    )
    make_table(doc,
        ["UI element", "what it actually is"],
        [
            ["Generate button",        "POST /api/generate which calls coordinator.generate which serially runs CLIP, UNet, VAE over httpx.AsyncClient."],
            ["Kill Pi VAE button",     "POST /api/admin/kill-vae which sends POST /admin/die to whichever worker has role=vae as its primary."],
            ["Per-stage timing rows",  "Coordinator's per-stage wall time. The via column is the worker that served that call. retry is how many transport-level failures happened before success."],
            ["Worker telemetry rows",  "Polled from each worker's /health endpoint every 2 s. RSS is psutil.Process().memory_info().rss inside that worker process."],
            ["Pulsing green alive dot","Last /health probe in the last 2 seconds came back 200."],
            ["Static red dead dot",    "Last /health probe failed with a transport error or non-200."],
            ["Fault inject dropdown",  "Tells the coordinator to schedule POST /admin/die on the chosen role 1 second after generate starts."],
        ]
    )

    # ===========================================================================
    # 11. TROUBLESHOOTING
    # ===========================================================================
    add_heading(doc, "11. Things that go wrong, and what to do", 1)

    add_heading(doc, "Worker won't start: 'Address already in use'", 2)
    doc.add_paragraph(
        "Another worker is still bound to that port. On Windows, find and kill it:"
    )
    code_block(doc,
        "Get-Process python | Where-Object { $_.MainWindowTitle -match 'worker' } | Stop-Process -Force\n"
        "# or, brute force\n"
        "Get-Process python | Stop-Process -Force"
    )
    doc.add_paragraph("On the Pi:")
    code_block(doc,
        "pkill -f 'worker.py --role vae'\n"
        "# or just hit Ctrl+C in the SSH session that owns it"
    )

    add_heading(doc, "Pi marked dead during generation", 2)
    doc.add_paragraph(
        "If the heartbeat monitor declares the Pi dead while the Pi is actually fine but "
        "busy, the issue is the worker pinning its event loop. We already fixed this by "
        "running torch in asyncio.to_thread. If it shows up again, double-check the Pi pulled "
        "the latest worker.py: cd ~/swarmgen and git pull."
    )

    add_heading(doc, "UI shows DEAD for a worker that is actually up", 2)
    doc.add_paragraph(
        "Two possibilities. First, the worker really did die: scroll back in its terminal and "
        "look for a Traceback. Second, the network path is blocked: from loq, "
    )
    code_block(doc, "curl http://<host>:<port>/health")
    doc.add_paragraph(
        "If curl works but the UI says dead, the API is calling the wrong URL. Check the "
        "--workers flag you passed to api.py."
    )

    add_heading(doc, "Pi runs out of memory mid VAE", 2)
    doc.add_paragraph(
        "Possible at 1.8 GB. The Pi will crash and the fallback on loq will pick up the next "
        "VAE call. If you want to keep the Pi alive for the demo, do not also poll /health "
        "at 100 ms in a tight loop while it is decoding. The eval harness uses 100 ms; the UI "
        "uses 2 seconds, which is fine."
    )

    add_heading(doc, "First generation is slow", 2)
    doc.add_paragraph(
        "First call after a worker boot runs the model lazily for the first time and triggers "
        "any GPU kernel JIT. Throw away the first run when reporting numbers. The eval and "
        "the baseline scripts already do this."
    )

    add_heading(doc, "mDNS hostname does not resolve", 2)
    doc.add_paragraph(
        "Skip mDNS and use the IP. We pass --workers explicitly to api.py for exactly this "
        "reason. If SSH cannot resolve raspberrypiinocula.local, use the IP 192.168.1.185 "
        "instead. mDNS is convenient when it works and a time sink when it does not."
    )

    # ===========================================================================
    # 12. TEAR DOWN
    # ===========================================================================
    add_heading(doc, "12. Shutting it down cleanly", 1)
    doc.add_paragraph(
        "Hit Ctrl+C in each of the three worker terminals and the api.py terminal. The "
        "FastAPI lifespan unregisters the mDNS announcement before exiting. Close the SSH "
        "sessions. Done."
    )

    # ===========================================================================
    # 13. SUBMISSION CHECKLIST
    # ===========================================================================
    add_heading(doc, "13. Submission checklist", 1)
    p = doc.add_paragraph(style="List Bullet"); p.add_run("Paper compiles on Overleaf with the IEEE conference template. Upload paper/paper.tex and the four PNGs in paper/figs.")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("Demo video uploaded as YouTube unlisted, 5 to 10 minutes, public link in the submission.")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("GitHub repo public for the duration of grading: github.com/sudarshan-sridhar/swarmgen.")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("Optional: flip the repo to private after submission.")

    # ===========================================================================
    # 14. ONE PAGE CHEAT SHEET
    # ===========================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix: one-page cheat sheet", 1)
    doc.add_paragraph(
        "Print this. Keep it next to you while recording."
    )

    add_heading(doc, "Hostnames and ports", 3)
    code_block(doc,
        "loq  192.168.1.16   port 8002   role unet (+ vae fallback)\n"
        "pc   192.168.1.39   port 8001   role clip\n"
        "pi   192.168.1.185  port 8003   role vae"
    )

    add_heading(doc, "Three terminals on loq", 3)
    code_block(doc,
        "# T1: loq (PowerShell, conda env ml)\n"
        "python worker.py --role unet --port 8002 --fallback-roles vae\n"
        "\n"
        "# T2: ssh into pc, then\n"
        "cd swarmgen && .venv\\Scripts\\activate.bat && python worker.py --role clip --port 8001\n"
        "\n"
        "# T3: ssh into pi, then\n"
        "cd ~/swarmgen && source .venv/bin/activate && python worker.py --role vae --port 8003\n"
        "\n"
        "# T4 (back on loq): API + UI\n"
        "python api.py --workers \"clip@192.168.1.39:8001,unet@192.168.1.16:8002,vae@192.168.1.185:8003\"\n"
        "# open http://localhost:7860"
    )

    add_heading(doc, "Numbers to mention", 3)
    p = doc.add_paragraph(style="List Bullet"); p.add_run("Single-device baseline: 380 ms / 2358 MB peak.")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("3-device with Pi VAE: 147 s / Pi 1226 MB peak (under 1845 MB).")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("3-device with loq VAE fallback: 1097 ms / 87.6 img/min batch.")
    p = doc.add_paragraph(style="List Bullet"); p.add_run("Fault recovery: about 250 ms from Pi kill to image complete.")

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size:,} bytes)")
